#AGGIUNGO LE LIBRERIE
import os
import logging
import time
import sys
import webbrowser
import string
import mechanize
import cookielib
import random
import re
from urllib2 import *
import socket
import getpass
import subprocess
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.shared import Inches
from docx.shared import Pt

#FUNZIONI DELLA LIBRERIA MECHANIZE DI PREPARAZIONE PER HONEYPOT
br= mechanize.Browser()
cj = cookielib.LWPCookieJar()
br.set_cookiejar(cj)
br.set_handle_equiv(True)
br.set_handle_redirect(True)
br.set_handle_referer(True)
br.set_handle_robots(False)
# Follows refresh 0 but not hangs on refresh > 0
br.set_handle_refresh(mechanize._http.HTTPRefreshProcessor(), max_time=1)
br.addheaders = [
    ('User-agent', 'Mozilla/5.0 (X11; U; Linux i686; en-US; rv:1.9.0.1) Gecko/2008071615 Fedora/3.0.1-1.fc9 Firefox/3.0.1')]

#FUNZIONE PER FAR AVVIARE I VARI COMANDI BASH
def bash_command(cmd):
    subprocess.Popen(cmd, shell=True, executable='/bin/bash')

#COLORI PER FUNZIONE HONEYPOT
info = '\033[1;33m[!]\033[1;m'
bad = '\033[1;31m[-]\033[1;m'
good = '\033[1;32m[+]\033[1;m'

#FUNZIONE DI CREAZIONE DEL FILE .docx INIZIALE
def creation_intro_file(file1,dominio1):
    doc=Document()
    report="---Report on web services---" + " " + dominio1
    title=doc.add_heading(report,0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titolo1="1-Analyse Vulnerabilities"
    paragraph = doc.add_heading().add_run(titolo1)
    font=paragraph.font
    font.color.rgb = RGBColor(0, 0, 205)
    doc.save(file1)

#5 FUNZIONI PER AGGIUNGERE I RISULTATI SUL FILE DOCX
def inserimento_figure(file1):
    f=open(file1,'rb')
    doc=Document(f)
    doc.add_picture('/home/oriol/Scrivania/PERSEUS_2.0/target/Screenshot_maps.png', width=Inches(7.20), height=Inches(4.20))
    doc.save(file1)

def aggiunta_titoloni(file1,aggiunta):
    f=open(file1,'rb')
    doc=Document(f)
    paragraph = doc.add_heading().add_run(aggiunta)
    font=paragraph.font
    font.color.rgb = RGBColor(0, 0, 205)
    doc.save(file1)

def aggiunta_comandi(file1,file2): 
    ff=open(file2,"r")
    Testo=ff.read()
    f=open(file1,'rb')
    doc=Document(f)
    inserimento=doc.add_paragraph(unicode(Testo,"utf-8"))
    doc.save(file1)

def aggiunta_altre_scritte_spazi(file1,testo):
    f=open(file1,'rb')
    doc=Document(f)
    inserimento=doc.add_paragraph(testo)
    doc.save(file1)

def aggiunta_titolo(file1,aggiunta):
    f=open(file1,'rb')
    doc=Document(f)
    inserimento=doc.add_paragraph(aggiunta)
    inserimento.alignment= WD_ALIGN_PARAGRAPH.CENTER
    doc.save(file1)
#FUNZIONE PER LEGGERE UN DOMINIO ALLA VOLTA
#FUZIONE PRINCIPALE 
def main():
    rosso="\033[31m"
    time.sleep(2)
    os.system("clear")
    #SCHERMATA INIZIALE E INSERENDO DOMINIO
    banner="bash" + " " + "/home/" + getpass.getuser() + "/Scrivania/PERSEUS_2.0/informazioni/banner.sh"
    os.system(banner)
    dominio = raw_input('\033[1;91mEnter Domain (e.g. google.com): \033[1;m')
    time.sleep(2)
    dom="echo" + " "+ dominio + " " + ">>" + " " + "dominio.txt"
    os.system(dom)
    permessi="chmod" + " " + "777" + " " + "/home/" + getpass.getuser() + "/Scrivania/PERSEUS_2.0/dominio.txt"
    os.system(permessi)
    
    dominio2=dominio.split(".")
    file=dominio2[0] + ".docx"
    creation_intro_file(file,dominio)

    red3=" "
    aggiunta_altre_scritte_spazi(file,red3)
    #CONTROLLO CHE IL SITO SIA ATTIVO E CHE NON CI SIA REVERSE PROXY ATTVO SUL SERVIZIO. IN CASO CONTRARIO VADO NELLA except
    aaa="++++++++1-2-CHECK DOMAIN and IP++++++++"
    aggiunta_titolo(file,aaa)
    ffff="provvisorio.txt"
    file33=open(ffff,"w")
    file33.close()
    x="curl -sLi --user-agent 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.31 (KHTML, like Gecko) Chrome/26.0.1410.63 Safari/537.31'" + " " + dominio + " " + "| grep -o 'HTTP/1.1 200 OK\|HTTP/2 200'" + " " + ">>" + " " + ffff
    os.system(x)
    aggiunta_comandi(file,ffff)
    try:
        addr=socket.gethostbyname(dominio)
    
    except socket.gaierror,err:
        elimina1="rm" + " " + "/home/" + getpass.getuser() + "/Scrivania/PERSEUS_2.0/" + file
        os.system(elimina1)
        delete2="rm" + " " + "/home/oriol/Scrivania/PERSEUS_2.0/provvisorio.txt"
        os.system(delete2)
        delete3="rm" + " " + "/home/oriol/Scrivania/PERSEUS_2.0/dominio.txt"
        os.system(delete3)
        eccezione(dominio,err)

    
    print ("\n")
    print("\033[1;31;34m 1-CHECK DOMAIN IF IS UP...DONE")
    print("\033[1;32;39m   ")
    
    x= "DOMINIO:" + " " + dominio + " " + " " + " " + " " + "IP:" + " " + addr
    aggiunta_altre_scritte_spazi(file,x)

    print("\033[1;31;34m 2-DOMAIN AND IP...DONE")
    print("\033[1;32;39m   ")
    red3=" "
    aggiunta_titolo(file,red3)

    e="++++++++3 GEO IP LOOKUP++++++++"
    aggiunta_titolo(file,e)
    geo_localizzazioneIP(dominio,ffff)
    aggiunta_comandi(file,ffff)
    inserimento_figure(file)
    remove="rm" + " " + "/home/" + getpass.getuser() + "/Scrivania/PERSEUS_2.0/target/Screenshot_maps.png"
    os.system(remove)
    print("\033[1;31;34m 3-GEO IP LOOKUP....DONE")
    print("\033[1;32;39m   ")
    
    
    red3=" "
    aggiunta_altre_scritte_spazi(file,red3)
   #whois da problemi ma non va nell except mettere condizione e farla per tutti gli altri
    try:
        a="++++++++4 WHOIS DOMAIN++++++++"
        aggiunta_titolo(file,a)
        red= "whois" + " " + dominio.encode(encoding='UTF-8',errors='strict') + " " + ">" + " " + ffff
    	os.system(red)
        time.sleep(1)
        aggiunta_comandi(file,ffff)
    

    except:
        print "EXCEPTION using the command WHOIS..."
        x= "EXCEPTION using the command WHOIS..."
        print sys.exc_info()[0]
        print sys.exc_info()[1]
        print sys.exc_info()[2]
        aggiunta_altre_scritte_spazi(file,x)
        time.sleep(5)
    
    print("\033[1;31;34m 4-WHOIS DOMAIN...DONE")
    print("\033[1;32;39m   ")
    red3=" "
    aggiunta_altre_scritte_spazi(file,red3)


    a="++++++++5 URL HISTORY WEB-SERVICES++++++++"
    aggiunta_titolo(file,a)
    storia='https://web.archive.org/web/*/' + dominio
    red2="echo" + " " + storia + " " + ">" + " " + ffff
    os.system(red2)
    time.sleep(1)
    aggiunta_comandi(file,ffff)
    
    print("\033[1;31;34m 5-HISTORY OF WEBSITE......DONE")
    print("\033[1;32;39m   ")
    red3=" "
    aggiunta_altre_scritte_spazi(file,red3)

    o="echo "+ " " + red + " " + ">" + " " + ffff
    os.system(o)

    delete2="rm" + " " + "/home/oriol/Scrivania/PERSEUS_2.0/provvisorio.txt"
    os.system(delete2)
    ffff="provvisorio.txt"
    file33=open(ffff,"w")
    file33.close()

    try:
        l="++++++++6 SPIDER SITE++++++++"
        aggiunta_titolo(file,l)
        spider="bash" + " " + "/home/oriol/Scrivania/PERSEUS_2.0/informazioni/spidershell.sh" + " " + ffff
        os.system(spider)
        time.sleep(1)
        aggiunta_comandi(file,ffff)

       
    
    except:
        print "TROUBLE USING A SPIDER ..."
        x= "TROUBLE USING A SPIDER ..."
        print sys.exc_info()[0]
        print sys.exc_info()[1]
        print sys.exc_info()[2]
        aggiunta_altre_scritte_spazi(file,x)
        time.sleep(5)
    
    print("\033[1;31;34m 6-SPIDER SITE....DONE")
    print("\033[1;32;39m   ")
    red3=" "
    aggiunta_altre_scritte_spazi(file,red3)   

    try:
        b="++++++++7 DNS LOOKUP++++++++"
        aggiunta_titolo(file,b)
        red= "host" + " " + dominio + " " + ">" + " " + ffff
        os.system(red)
        aggiunta_comandi(file,ffff)
    	
    
    except:
        print "ERROR DETECT DNS ..."
        x= "ERROR DETECT DNS ..."
        print sys.exc_info()[0]
        print sys.exc_info()[1]
        print sys.exc_info()[2]
        aggiunta_altre_scritte_spazi(file,x)
        time.sleep(5)

    print("\033[1;31;34m 7-DNS LOOKUP........DONE")
    print("\033[1;32;39m   ")
    red3=" "
    aggiunta_altre_scritte_spazi(file,red3)   
    
    try:
        gg="++++++++8 DETECT HONEYPOT++++++++"
        aggiunta_titolo(file,gg)
        honeypot(addr,ffff,file)
        time.sleep(1)
        remove="rm" + " " + "/home/" + getpass.getuser() + "/Scrivania/PERSEUS_2.0/provvisorio1.txt"
        bash_command(remove)
        time.sleep(1)
        print("\033[1;31;34m 8-DETECT HONEYPOT...DONE")
        print("\033[1;32;39m   ")
        red3=" "
        aggiunta_altre_scritte_spazi(file,red3)
    
    except:
        print "ERROR DETECT HONEYPOT ..."
        x= "ERROR DETECT HONEYPOT ..."
        print sys.exc_info()[0]
        print sys.exc_info()[1]
        print sys.exc_info()[2]
        aggiunta_altre_scritte_spazi(file,x)
        time.sleep(3)
 
    gggg="++++++++9 SUB-DOMAINS++++++++"
    aggiunta_titolo(file,gggg)
    red=" "
    o="echo "+ " " + red + " " + ">" + " " + ffff
    bash_command(o)
    path="/home/oriol/Scrivania/PERSEUS_2.0/informazioni/sub.py"
    sub="python" + " " + path + " " + dominio + " " + "-o" + " " + ffff
    os.system(sub)
    aggiunta_comandi(file,ffff)

    print("\033[1;31;34m 9-SUB-DOMAINS.........DONE")
    print("\033[1;32;39m   ")
    red=" "
    o="echo "+ " " + red + " " + ">" + " " + ffff
    os.system(o)
    red3=" "
    aggiunta_altre_scritte_spazi(file,red3)
    
    try:
        lo="++++++++10 CMS DETECTION++++++++"
        aggiunta_titolo(file,lo)
        red=" "
        o="echo "+ " " + red + " " + ">" + " " + ffff
        os.system(o)
        time.sleep(1)
        avvio="python" + " " + "/home/" + getpass.getuser() + "/Scrivania/PERSEUS_2.0/informazioni/cms.py" + " " + dominio + " " + ffff
        os.system(avvio)
        aggiunta_comandi(file,ffff)
        print("\033[1;31;34m 10-CMS DETECTION........DONE")
        print("\033[1;32;39m   ")
    except:
        print "ERROR DETECT CMS ..."
        x= "ERROR DETECT CMS ..."
        print sys.exc_info()[0]
        print sys.exc_info()[1]
        print sys.exc_info()[2]
        aggiunta_altre_scritte_spazi(file,x)
        time.sleep(3)
    
    
    red3=" "
    aggiunta_altre_scritte_spazi(file,red3)
    red=" "
    o="echo "+ " " + red + " " + ">" + " " + ffff
    os.system(o)

    lol="++++++++11 SERVER WEB VERSION DETECTION++++++++"
    aggiunta_titolo(file,lol)
    red3="proxychains"+ " "+ "HEAD" + " " + dominio + " " + ">" + " " + ffff
    os.system(red3)
    aggiunta_comandi(file,ffff)
    print("\033[1;31;34m 11-SERVER WEB VERSION DETECT.....DONE")
    print("\033[1;32;39m   ")
    
    red=" "
    o="echo "+ " " + red + " " + ">" + " " + ffff
    os.system(o)
    red3=" "
    aggiunta_altre_scritte_spazi(file,red3)

    h1="++++++++12 ZONE TRANSFER DNS++++++++"
    aggiunta_titolo(file,h1)
    zon = "http://api.hackertarget.com/zonetransfer/?q=" + dominio
    try:
        tran = urlopen(zon).read()
        out_file = open(ffff,"w")
        out_file.write(tran)
        out_file.close()
        aggiunta_comandi(file,ffff)
    except Exception as inst:
        error=str(inst)
        red44="echo" + " " + error + " "+ ">" + " " + ffff
        os.system(red44)
        time.sleep(1)
        aggiunta_comandi(file,ffff)
        pass

    print("\033[1;31;34m 12-ZONE TRANSFER DNS....DONE")
    print("\033[1;32;39m   ")
    red=" "
    o="echo "+ " " + red + " " + ">" + " " + ffff
    os.system(o)
    red3=" "
    aggiunta_altre_scritte_spazi(file,red3)

    h2="++++++++13 SUBNET LOOKUP++++++++"
    aggiunta_titolo(file,h2)
    sub = "http://api.hackertarget.com/subnetcalc/?q=" + dominio
    try:
        net = urlopen(sub).read()
        out_file = open(ffff,"w")
        out_file.write(net)
        out_file.close()
        aggiunta_comandi(file,ffff)
    except Exception as inst:
        error=str(inst)
        red44="echo" + " " + error + " "+ ">" + " " + ffff
        os.system(red44)
        aggiunta_comandi(file,ffff)
        pass

    print("\033[1;31;34m 13-SUBNET LOOKUP.......DONE")
    print("\033[1;32;39m   ")
    red=" "
    o="echo "+ " " + red + " " + ">" + " " + ffff
    os.system(o)
    red3=" "
    aggiunta_altre_scritte_spazi(file,red3)

    tr="++++++++14 TRACEROUTE++++++++"
    aggiunta_titolo(file,tr)
    red= "proxychains" + " " + "traceroute" + " " + dominio + " " + ">" + " " + ffff
    os.system(red)
    aggiunta_comandi(file,ffff)
    print("\033[1;31;34m 14-TRACEROUTE.........DONE")
    print("\033[1;32;39m   ")
    
    red=" "
    o="echo "+ " " + red + " " + ">" + " " + ffff
    os.system(o)
    red3=" "
    aggiunta_altre_scritte_spazi(file,red3)
    
    f="++++++++15 PORT DETECTION++++++++"
    aggiunta_titolo(file,f)
    command2="echo" + " " + "orologio96" + " " + "|" + " " + "sudo" + " " + "-S" + " " + "proxychains" + " " + "nmap" +  " " + "-sT" + " " +  "-PN" + " " + "-n" + " " + "-sV" + " " + "-p" + " " + "21,22,53,80,110,139,143,443,3306" + " " + addr + " " + ">" + " " + ffff
    os.system(command2)
    copia="cat" + " " + ffff + " " + ">" + " " + "provvisorio1.txt"
    os.system(copia)
    aggiunta_comandi(file,ffff)
    print("\033[1;31;34m 15-PORT DETECTION......DONE")
    print("\033[1;32;39m   ")

    red=" "
    o="echo "+ " " + red + " " + ">" + " " + ffff
    os.system(o)
    red3=" "
    aggiunta_altre_scritte_spazi(file,red3)   
    
    h="++++++++16 OPERATING SYSTEM OF TARGET MACHINE++++++++"
    aggiunta_titolo(file,h)
    command2="echo" + " " + "orologio96" + " " + "|" + " " + "sudo" + " " + "-S" + " " + "proxychains" + " " + "nmap" +  " " + "-O" + " " + addr + " " + ">" + " " + ffff
    os.system(command2)
    aggiunta_comandi(file,ffff)
    print("\033[1;31;34m 16-OPERATING SYSTEM OF TARGET MACHINE....DONE")
    print("\033[1;32;39m   ")
    red=" "
    o="echo "+ " " + red + " " + ">" + " " + ffff
    os.system(o)
    red3=" "
    aggiunta_altre_scritte_spazi(file,red3)
  
    h1="++++++++17 EXTRACT EMAIL FROM WEB-SITE++++++++"
    aggiunta_titolo(file,h1)
    email="echo" + " " + "orologio96" + " " + "|" + " " + "sudo" + " " + "-S" + " " + "bash" + " " + "/home/oriol/Scrivania/PERSEUS_2.0/informazioni/email.sh" + " " + dominio + " " + ffff
    os.system(email)
    aggiunta_comandi(file,ffff)
    print("\033[1;31;34m 17-EXTRACT EMAIL FROM WEB-SITE....DONE")
    print("\033[1;32;39m   ")
    
    red=" "
    o="echo "+ " " + red + " " + ">" + " " + ffff
    os.system(o)
    red3=" "
    aggiunta_altre_scritte_spazi(file,red3)
    remove="echo" + " " + "orologio96" + " " + "|" + " " + "sudo" + " " + "-S" + " " + "bash" + " " + "/home/oriol/Scrivania/PERSEUS_2.0/delete.sh"
    bash_command(remove) 

    f1="++++++++18 DBMS DETECTION ON PORT 3306++++++++"
    aggiunta_titolo(file,f1)
    fff="provvisorio1.txt"
    copia2="cp" + " " + fff + " " + " /home/oriol/Scrivania/PERSEUS_2.0/informazioni"
    os.system(copia2)
    detect= "echo" + " " + "orologio96" + " " + "|" + " " + "sudo" + " " + "-S" + " " "bash" + " " + "/home/oriol/Scrivania/PERSEUS_2.0/informazioni/dbms1.sh" + " " + addr + " " + fff
    os.system(detect)
    aggiunta_comandi(file,fff)
    print("\033[1;31;34m 18-DBMS DETECTION ON PORT 3306....DONE")
    print("\033[1;32;39m   ")
    
    f1233="++++++++19 SUMMARY TECHNOLOGIES DETECTION++++++++"
    aggiunta_titolo(file,f1233)
    try:
        command="wad" + " " + "-u" + " " + "https://www." + dominio + " " + ">" + " " + ffff
        os.system(command)
        aggiunta_comandi(file,ffff)

    except ssl.SSLError as err:
        errore="SSL connection failed. I'm not able to run wad on " + " " + dominio + " " + ">" + " " + ffff
        aggiunta_comandi(file,ffff) 

    print("\033[1;31;34m 19-SUMMARY TECHNOLOGIES DETECTION....DONE")
    print("\033[1;32;39m   ")


    ricerca_file()
   
    delete2="rm" + " " + "/home/oriol/Scrivania/PERSEUS_2.0/provvisorio.txt"
    delete3="rm" + " " + "/home/oriol/Scrivania/PERSEUS_2.0/dominio.txt"
    delete4="rm" + " " + "/home/oriol/Scrivania/PERSEUS_2.0/provvisorio1.txt"
    os.system(delete2)
    os.system(delete3)
    os.system(delete4)
    copia2="cp" + " " + file + " " + " /home/oriol/Scrivania/PERSEUS_2.0/target"
    os.system(copia2)
    delete1="rm" + " " + "/home/oriol/Scrivania/PERSEUS_2.0/" + file
    os.system(delete1)

    print "END TO PICK UP INFO OF THE WEB SERVICES"
    time.sleep(5)

#RICERCARE FILE DA DENTRO UNA CARTELLA E CANCELLARE QUELLI SELEZIONATI
def ricerca_file():
    dir='/home/oriol/Scrivania/PERSEUS_2.0/informazioni'
    files=os.listdir(dir)
    for file1 in files:
        if file1=="provvisorio1.txt":
            delete2="rm" + " " + "/home/oriol/Scrivania/PERSEUS_2.0/informazioni/provvisorio1.txt"
            os.system(delete2)

        if file1=="meterpreter.rc":
            print "FILE esiste"
            remove= "sudo" + " " + "bash" + " " + "/home/oriol/Scrivania/PERSEUS_2.0/informazioni/delete1.sh"
            bash_command(remove)
        



#QUATTRO FUNZIONI PER LA GEOLOCALIZZAZIONE
def geo_localizzazioneIP(dominio1,file1): 
    bot_sulla_pagina(dominio1,file1)


def ottengo_IP(dominio2):
    addr=socket.gethostbyname(dominio2)
    add1=str(addr)
    return add1

def bot_sulla_pagina(dominio1,file2):
    driver = webdriver.Firefox(executable_path=r'/home/oriol/Scrivania/PERSEUS_2.0/informazioni/geckodriver')
    stringa="https://www.maxmind.com/en/geoip2-precision-demo"
    driver.get(stringa)
    address=ottengo_IP(dominio1)
    elem=driver.find_element_by_xpath("//form[@id='geoip-demo-form']/textarea")
    elem.send_keys(address)
    button=driver.find_element_by_xpath("//form[@id='geoip-demo-form']/button")
    button.click()
    time.sleep(1)
    risultati=driver.find_element_by_xpath("//div[@id='geoip-demo']/div[@class='table-overflow-container']/table[@class='table table-bordered table-striped table-condensed table-small']/tbody[@id='geoip-demo-results-tbody']/tr/td[5]")#/tr[@class='geoip-results']") 
    coordinate=risultati.text
    stringa=string.split(coordinate)
    red2="echo" + " " + " Latitude: " + " " + stringa[0] + " " + ">" + " " + file2
    red3="echo" + " " + " Longitude: " + " " + stringa[1] + " " + ">>" + " " + file2
    bash_command(red2)
    bash_command(red3)
    photo_maps(coordinate)
    driver.quit()


def photo_maps(coordinate1):
    driver = webdriver.Firefox(executable_path=r'/home/oriol/Scrivania/PERSEUS_2.0/informazioni/geckodriver')  
    maps='https://www.google.com/maps/place/+' + coordinate1
    driver.get(maps)
    name1="Screenshot_maps.png"
    driver.maximize_window()
    screenshot = driver.save_screenshot(name1)
    time.sleep(3)
    copia2="cp" + " " + name1 + " " + "/home/" + getpass.getuser() + "/Scrivania/PERSEUS_2.0/target"
    os.system(copia2)
    remove="rm" + " " + "/home/" + getpass.getuser() + "/Scrivania/PERSEUS_2.0/" + name1
    os.system(remove)
    driver.quit()


#CREO LA FUNZIONE ECCEZIONE SE IL DOMINIO e' DOWN OPPURE PROTETTO DA REVERSE PROXY
def eccezione(dominio1,err1):
    print "Cannot resolve hostname....: ", dominio1, err1
    print "\033[1;31;34m It could be protect by Reverse Proxy like Cloudflare\n"
    time.sleep(2)
    dominio_intero="www."+ dominio1
    #QUA NON POSSO USARE getpass.getuser()!!!!!!!!!!!!
    driver = webdriver.Firefox(executable_path=r'/home/oriol/Scrivania/PERSEUS_2.0/informazioni/geckodriver')
    stringa="https://hostingchecker.com/"
    driver.get(stringa)
    elem=driver.find_element_by_tag_name("input")
    elem.send_keys(dominio_intero)
    elem.send_keys(Keys.ENTER)
    time.sleep(4)
    name1="trouble"  + ".png"
    driver.maximize_window()
    screenshot = driver.save_screenshot(name1)
    document=Document()
    report="Report on web services--->" + " " + dominio1
    title=document.add_heading().add_run(report,0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font=title.font
    font.color.rgb = RGBColor(0, 0, 205)
    figura=document.add_picture(name1, width=Inches(6.50), height=Inches(5.20))
    figura.alignment = WD_ALIGN_PARAGRAPH.LEFT

    stringa="From here I can tell you that being the service hosted by a Reverse Proxy, it can not be analyzed unless you get the real IP address"
    p=document.add_paragraph().add_run(stringa)
    font=p.font
    font.color.rgb = RGBColor(255, 0, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    dominio2=dominio1.split('.')
    name=dominio2[0] + ".docx"
    document.save(name)
    elimina1="rm" + " " + "/home/" + getpass.getuser() + "/Scrivania/PERSEUS_2.0/" + name1
    os.system(elimina1)
    copia2="cp" + " " + name + " " + " /home/oriol/Scrivania/PERSEUS_2.0/target"
    os.system(copia2)
    elimina2="rm" + " " + "/home/" + getpass.getuser() + "/Scrivania/PERSEUS_2.0/" + name
    os.system(elimina2)
    driver.quit()
    sys.exit()

#FUNZIONE PER RILEVAZIONE DI HONEYPOT SUL WEB-SERVICE ANALIZZATO
def honeypot(addr,file,file1):
    result = {"0.0": 0, "0.1": 10, "0.2": 20, "0.3": 30, "0.4": 40, "0.5": 50, "0.6": 60, "0.7": 70, "0.8": 80, "0.9": 90, "1.0": 10}
    honey = 'https://api.shodan.io/labs/honeyscore/%s?key=C23OXE0bVMrul2YeqcL7zxb6jZ4pj2by' % addr
    file44="provvisorio1.txt"
    try:
        phoney = br.open(honey).read()
        if float(phoney) >= 0.0 and float(phoney) <= 0.4:
            what = good
        else:
            what = bad
        honeys= '{} Honeypot Probabilty: {}%'.format(what, result[phoney])
        c= "echo"+ " " + honeys
        out_file = open(file,"w")
        out_file.write(honeys)
    	out_file.close()
        command="sed" + " " + "'s/................//' " + " " + file + ">" + " " +  file44
        bash_command(command)
        time.sleep(1)
        command2="cat" + " " + file44 + " " ">" + " " + file
        bash_command(command2)
        time.sleep(1)
        aggiunta_comandi(file1,file)

    except KeyError:
        honeys=' Honeypot prediction failed'
        out_file = open(file,"w")
        out_file.write(honeys)
    	out_file.close()
        command="sed" + " " + "'s/................//' " + " " + file + ">" + " " +  file44
        bash_command(command)
        time.sleep(1)
        command2="cat" + " " + file44 + " " ">" + " " + file
        bash_command(command2)
        time.sleep(1)
        aggiunta_comandi(file1,file)

if __name__ == '__main__':
    main()
